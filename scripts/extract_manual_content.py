from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document

from extract_chapter_content import (
    CURATED_TERMS,
    DEFAULT_EN_DOCX,
    DEFAULT_REPO_ROOT,
    DEFAULT_ZH_DOCX,
    attach_assets_to_lesson,
    build_lesson,
    iter_doc_items,
    normalize_heading_item,
    normalize_title,
    serialize_block,
    write_json,
)


CN_NUMERALS = [
    "一",
    "二",
    "三",
    "四",
    "五",
    "六",
    "七",
    "八",
    "九",
    "十",
    "十一",
    "十二",
    "十三",
    "十四",
    "十五",
    "十六",
    "十七",
    "十八",
    "十九",
    "二十",
    "二十一",
    "二十二",
    "二十三",
    "二十四",
    "二十五",
    "二十六",
    "二十七",
    "二十八",
    "二十九",
    "三十",
    "三十一",
    "三十二",
    "三十三",
]


def cell_text(row: Any, index: int) -> str:
    return " ".join(row.cells[index].text.split()) if len(row.cells) > index else ""


def chapter_number(label: str) -> int | None:
    match = re.search(r"(\d+)", label)
    if match:
        return int(match.group(1))
    return None


def parse_toc(docx_path: Path, lang: str) -> list[dict[str, Any]]:
    doc = Document(str(docx_path))
    chapters: list[dict[str, Any]] = []
    for row in doc.tables[0].rows:
        label = cell_text(row, 0)
        title = cell_text(row, 1)
        page_text = cell_text(row, 2)
        if lang == "en" and not label.startswith("Chapter "):
            continue
        if lang == "zh" and not label.startswith("第 "):
            continue
        number = chapter_number(label)
        if not number:
            continue
        chapters.append(
            {
                "chapter": number,
                "label": label,
                "title": title,
                "pageStart": int(page_text),
            }
        )
    return chapters


def find_starts(items: list[dict[str, Any]], chapters: list[dict[str, Any]], lang: str) -> dict[int, int]:
    starts: dict[int, int] = {}
    for chapter in chapters:
        number = chapter["chapter"]
        title = chapter["title"]
        for index, item in enumerate(items):
            text = item.get("text", "")
            style = item.get("style", "")
            if item.get("kind") != "paragraph" or not style.startswith("Heading"):
                continue
            if lang == "en" and f"Chapter {number}:" in text:
                starts[number] = index
                break
            if lang == "zh":
                if "内部目录" in text:
                    continue
                cn = CN_NUMERALS[number - 1]
                prefixes = (f"第 {number} 章", f"第{number}章", f"第{cn}章")
                if any(prefix in text for prefix in prefixes):
                    starts[number] = index
                    break
        if number not in starts:
            raise RuntimeError(f"Could not find {lang} start for chapter {number}: {title}")
    return starts


def normalize_items_for_chapter(items: list[dict[str, Any]], lang: str) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    for item in items:
        normalized.extend(normalize_heading_item(item, lang))
    return normalized


def slugify(text: str) -> str:
    normalized = normalize_title(text)
    slug = re.sub(r"[^a-z0-9]+", "-", normalized).strip("-")
    return slug or "section"


def unique_section_id(chapter_id: str, used: set[str], index: int, title: str) -> str:
    base = f"{chapter_id}-s{index:02d}-{slugify(title)}"
    candidate = base
    suffix = 2
    while candidate in used:
        candidate = f"{base}-{suffix}"
        suffix += 1
    used.add(candidate)
    return candidate


def load_source_toc(repo_root: Path) -> dict[int, dict[str, Any]]:
    toc_path = repo_root / "content" / "source" / "source_toc_sections.json"
    if not toc_path.exists():
        return {}
    data = json.loads(toc_path.read_text(encoding="utf-8"))
    return {int(chapter["chapter"]): chapter for chapter in data.get("chapters", [])}


def mapped_source_page(
    source_page: int,
    source_start: int,
    source_next: int,
    target_start: int,
    target_next: int,
) -> int:
    if source_next <= source_start or target_next <= target_start:
        return target_start
    ratio = (source_page - source_start) / (source_next - source_start)
    mapped = target_start + int((target_next - target_start) * ratio + 0.5)
    return max(target_start, min(target_next - 1, mapped))


def heading_candidates(items: list[dict[str, Any]]) -> list[tuple[int, dict[str, Any]]]:
    return [
        (index, item)
        for index, item in enumerate(items)
        if item.get("kind") == "heading" and item.get("text")
    ]


def aligned_heading_position(
    heading_ordinal: int,
    source_item_index: int,
    source_item_count: int,
    target_item_count: int,
    target_headings: list[tuple[int, dict[str, Any]]],
    previous_position: int,
) -> int | None:
    if not target_headings:
        return None

    estimated_position = int((source_item_index / max(1, source_item_count)) * target_item_count + 0.5)
    ordinal_window = range(max(0, heading_ordinal - 2), min(len(target_headings), heading_ordinal + 3))
    candidates = [
        (ordinal, target_headings[ordinal][0])
        for ordinal in ordinal_window
        if target_headings[ordinal][0] > previous_position
    ]
    if not candidates:
        candidates = [
            (ordinal, position)
            for ordinal, (position, _item) in enumerate(target_headings)
            if position > previous_position
        ]
    if not candidates:
        return None

    best_score = min(abs(position - estimated_position) for _ordinal, position in candidates)
    max_reasonable_distance = max(8, int(target_item_count * 0.05))
    use_preferred_ordinal = True
    if best_score > max_reasonable_distance:
        global_candidates = [
            (ordinal, position)
            for ordinal, (position, _item) in enumerate(target_headings)
            if position > previous_position
        ]
        if global_candidates:
            candidates = global_candidates
            best_score = min(abs(position - estimated_position) for _ordinal, position in candidates)
            use_preferred_ordinal = False

    preferred = next((position for ordinal, position in candidates if ordinal == heading_ordinal), None)
    if use_preferred_ordinal and preferred is not None and abs(preferred - estimated_position) <= best_score + 4:
        tied_after = [
            position
            for _ordinal, position in candidates
            if abs(position - estimated_position) == abs(preferred - estimated_position)
            and position >= estimated_position
        ]
        return min(tied_after) if tied_after else preferred

    return min(
        candidates,
        key=lambda candidate: (
            abs(candidate[1] - estimated_position),
            0 if candidate[1] >= estimated_position else 1,
            candidate[1],
        ),
    )[1]


def title_matches(candidate: str, expected: str) -> bool:
    candidate_norm = normalize_title(candidate)
    expected_norm = normalize_title(expected)
    if not candidate_norm or not expected_norm:
        return False
    if candidate_norm == expected_norm:
        return True
    return candidate_norm.startswith(expected_norm) or expected_norm.startswith(candidate_norm)


def find_source_section_positions(
    items: list[dict[str, Any]],
    source_sections: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    headings = heading_candidates(items)
    heading_ordinal_by_index = {
        item_index: ordinal
        for ordinal, (item_index, _item) in enumerate(heading_candidates(items))
    }
    matches: list[dict[str, Any]] = []
    last_item_index = -1
    last_heading_ordinal = -1
    for source_section in source_sections:
        found: tuple[int, int, dict[str, Any]] | None = None
        for ordinal, (item_index, item) in enumerate(headings):
            if item_index <= last_item_index or ordinal <= last_heading_ordinal:
                continue
            if title_matches(item.get("text", ""), source_section["title"]):
                found = (ordinal, item_index, item)
                break
        if found is None:
            continue
        ordinal, item_index, item = found
        matches.append(
            {
                "source": source_section,
                "itemIndex": item_index,
                "headingOrdinal": heading_ordinal_by_index.get(item_index, ordinal),
                "matchKind": item.get("kind"),
                "title": item.get("text", source_section["title"]),
            }
        )
        last_item_index = item_index
        last_heading_ordinal = ordinal
    return matches


def content_has_blocks(items: list[dict[str, Any]]) -> bool:
    return any(item.get("kind") != "heading" or item.get("text") for item in items)


def serialize_items(section_id: str, lang: str, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        serialize_block(item, section_id, lang, index)
        for index, item in enumerate(items, start=1)
    ]


def build_sectionized_chapter(
    chapter: int,
    page_start: int,
    page_end: int,
    en_title: str,
    zh_title: str,
    en_body: list[dict[str, Any]],
    zh_body: list[dict[str, Any]],
    source_chapter: dict[str, Any],
    next_source_page: int,
    next_target_page: int,
) -> list[dict[str, Any]]:
    chapter_id = f"ch{chapter:02d}"
    matches = find_source_section_positions(en_body, source_chapter.get("sections", []))
    if not matches:
        section_id = f"{chapter_id}-content"
        return [
            {
                "id": section_id,
                "level": 1,
                "page": page_start,
                "title": {"en": en_title, "zh": zh_title},
                "content": {
                    "en": serialize_items(section_id, "en", en_body),
                    "zh": serialize_items(section_id, "zh", zh_body),
                },
            }
        ]

    zh_heading_candidates = heading_candidates(zh_body)
    used_ids: set[str] = set()
    sections: list[dict[str, Any]] = []
    zh_positions: list[int | None] = []
    previous_zh_position = -1
    for match in matches:
        zh_position = aligned_heading_position(
            match["headingOrdinal"],
            match["itemIndex"],
            len(en_body),
            len(zh_body),
            zh_heading_candidates,
            previous_zh_position,
        )
        zh_positions.append(zh_position)
        if zh_position is not None:
            previous_zh_position = zh_position

    en_ranges: list[dict[str, Any]] = []
    first_match_index = matches[0]["itemIndex"]
    if first_match_index > 0 and content_has_blocks(en_body[:first_match_index]):
        intro_title = en_title
        first_source_section = source_chapter.get("sections", [None])[0]
        if first_source_section and first_source_section.get("title", "").lower().startswith("what is "):
            intro_title = first_source_section["title"]
        en_ranges.append(
            {
                "title": intro_title,
                "page": page_start,
                "level": 1,
                "enStart": 0,
                "enEnd": first_match_index,
                "zhStart": 0,
                "zhEnd": None,
                "zhTitle": zh_title,
            }
        )

    for index, match in enumerate(matches):
        next_match_index = matches[index + 1]["itemIndex"] if index + 1 < len(matches) else len(en_body)
        zh_position = zh_positions[index]
        next_zh_position = zh_positions[index + 1] if index + 1 < len(zh_positions) else None
        source_page = int(match["source"]["sourcePage"])
        page = mapped_source_page(
            source_page,
            int(source_chapter["sourcePage"]),
            next_source_page,
            page_start,
            next_target_page,
        )
        zh_title_for_section = (
            zh_body[zh_position].get("text", match["title"])
            if zh_position is not None and zh_position < len(zh_body)
            else match["title"]
        )
        en_ranges.append(
            {
                "title": match["title"],
                "page": page,
                "level": 2,
                "enStart": match["itemIndex"] + 1,
                "enEnd": next_match_index,
                "zhStart": zh_position + 1 if zh_position is not None else None,
                "zhEnd": next_zh_position,
                "zhTitle": zh_title_for_section,
            }
        )

    for index, entry in enumerate(en_ranges):
        if entry["zhEnd"] is None:
            next_known_zh = next((later["zhStart"] for later in en_ranges[index + 1:] if later["zhStart"] is not None), None)
            entry["zhEnd"] = next_known_zh if next_known_zh is not None else len(zh_body)
        if entry["zhStart"] is None:
            prev_zh_end = en_ranges[index - 1]["zhEnd"] if index > 0 else 0
            entry["zhStart"] = prev_zh_end

    for index, entry in enumerate(en_ranges, start=1):
        en_segment = en_body[entry["enStart"]:entry["enEnd"]]
        zh_segment = zh_body[entry["zhStart"]:entry["zhEnd"]]
        if not content_has_blocks(en_segment) and not content_has_blocks(zh_segment):
            continue
        section_id = unique_section_id(chapter_id, used_ids, index, entry["title"])
        sections.append(
            {
                "id": section_id,
                "level": entry["level"],
                "page": entry["page"],
                "title": {
                    "en": entry["title"],
                    "zh": entry["zhTitle"],
                },
                "content": {
                    "en": serialize_items(section_id, "en", en_segment),
                    "zh": serialize_items(section_id, "zh", zh_segment),
                },
            }
        )

    return sections


def generic_chapter(
    chapter: int,
    page_start: int,
    page_end: int,
    en_title: str,
    zh_title: str,
    en_items: list[dict[str, Any]],
    zh_items: list[dict[str, Any]],
    en_docx: Path,
    zh_docx: Path,
    source_chapter: dict[str, Any] | None = None,
    next_source_page: int | None = None,
    next_target_page: int | None = None,
) -> dict[str, Any]:
    chapter_id = f"ch{chapter:02d}"
    section_id = f"{chapter_id}-content"
    en_body = normalize_items_for_chapter(en_items[1:], "en")
    zh_body = normalize_items_for_chapter(zh_items[1:], "zh")
    if source_chapter and next_source_page and next_target_page:
        sections = build_sectionized_chapter(
            chapter,
            page_start,
            page_end,
            en_title,
            zh_title,
            en_body,
            zh_body,
            source_chapter,
            next_source_page,
            next_target_page,
        )
        extraction = "python-docx body-order extraction; source-TOC sectionized bilingual content"
    else:
        en_blocks = [serialize_block(item, section_id, "en", idx) for idx, item in enumerate(en_body, start=1)]
        zh_blocks = [serialize_block(item, section_id, "zh", idx) for idx, item in enumerate(zh_body, start=1)]
        sections = [
            {
                "id": section_id,
                "level": 1,
                "page": page_start,
                "title": {
                    "en": en_title,
                    "zh": zh_title,
                },
                "content": {
                    "en": en_blocks,
                    "zh": zh_blocks,
                },
            }
        ]
        extraction = "python-docx body-order extraction; generic chapter-level bilingual content"
    return attach_assets_to_lesson({
        "id": chapter_id,
        "chapter": chapter,
        "pageStart": page_start,
        "pageEnd": page_end,
        "title": {
            "en": f"Chapter {chapter}: {en_title}",
            "zh": f"第 {chapter} 章：{zh_title}",
        },
        "source": {
            "enDocx": str(en_docx),
            "zhDocx": str(zh_docx),
            "extraction": extraction,
        },
        "sections": sections,
    })


def collect_referenced_assets(manual: dict[str, Any]) -> dict[str, str]:
    assets: dict[str, str] = {}
    for chapter in manual["chapters"]:
        for asset in chapter.get("assets", []):
            assets[asset["id"]] = asset["path"]
    return assets


def read_docx_media(docx_path: Path) -> dict[str, bytes]:
    assets: dict[str, bytes] = {}
    with ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/"):
                continue
            blob = archive.read(name)
            digest = hashlib.sha256(blob).hexdigest()
            asset_id = f"fig-{digest[:16]}"
            assets.setdefault(asset_id, blob)
    return assets


def write_asset_files(repo_root: Path, manual: dict[str, Any], en_docx: Path, zh_docx: Path) -> None:
    referenced_assets = collect_referenced_assets(manual)
    public_content_dir = repo_root / "apps" / "reader" / "public" / "content"
    public_figures_dir = public_content_dir / "assets" / "figures"
    processed_figures_dir = repo_root / "content" / "processed" / "assets" / "figures"
    for directory in [public_figures_dir, processed_figures_dir]:
        if directory.exists():
            shutil.rmtree(directory)
        directory.mkdir(parents=True, exist_ok=True)

    media_assets = read_docx_media(en_docx)
    media_assets.update(read_docx_media(zh_docx))
    missing = sorted(asset_id for asset_id in referenced_assets if asset_id not in media_assets)
    if missing:
        raise RuntimeError(f"Missing media blobs for referenced assets: {', '.join(missing[:5])}")

    manifest_assets: list[dict[str, Any]] = []
    for asset_id, relative_path in sorted(referenced_assets.items()):
        blob = media_assets[asset_id]
        output_name = Path(relative_path).name
        for directory in [public_figures_dir, processed_figures_dir]:
            (directory / output_name).write_bytes(blob)
        manifest_assets.append({"id": asset_id, "path": relative_path, "bytes": len(blob)})

    write_json(
        public_content_dir / "assets" / "asset-manifest.json",
        {
            "version": manual["version"],
            "assetCount": len(manifest_assets),
            "assets": manifest_assets,
        },
    )


def build_manual(en_docx: Path, zh_docx: Path, repo_root: Path = DEFAULT_REPO_ROOT) -> dict[str, Any]:
    en_chapters = parse_toc(en_docx, "en")
    zh_chapters = parse_toc(zh_docx, "zh")
    if len(en_chapters) != len(zh_chapters):
        raise RuntimeError(f"Chapter TOC mismatch: en={len(en_chapters)}, zh={len(zh_chapters)}")
    source_toc = load_source_toc(repo_root)

    en_items = iter_doc_items(Document(str(en_docx)))
    zh_items = iter_doc_items(Document(str(zh_docx)))
    en_starts = find_starts(en_items, en_chapters, "en")
    zh_starts = find_starts(zh_items, zh_chapters, "zh")

    lessons: list[dict[str, Any]] = []
    for index, (en_chapter, zh_chapter) in enumerate(zip(en_chapters, zh_chapters)):
        chapter = en_chapter["chapter"]
        if chapter != zh_chapter["chapter"]:
            raise RuntimeError(f"Chapter number mismatch: {chapter} vs {zh_chapter['chapter']}")
        page_start = int(en_chapter["pageStart"])
        next_page = int(en_chapters[index + 1]["pageStart"]) if index + 1 < len(en_chapters) else 450
        page_end = next_page - 1
        if chapter == 1:
            lesson = build_lesson(en_docx, zh_docx)
        else:
            en_start = en_starts[chapter]
            zh_start = zh_starts[chapter]
            next_chapter = chapter + 1
            en_end = en_starts.get(next_chapter, len(en_items))
            zh_end = zh_starts.get(next_chapter, len(zh_items))
            source_chapter = source_toc.get(chapter)
            next_source_page = None
            if source_chapter:
                next_source = source_toc.get(next_chapter)
                next_source_page = int(next_source["sourcePage"]) if next_source else 558
            lesson = generic_chapter(
                chapter,
                page_start,
                page_end,
                en_chapter["title"],
                zh_chapter["title"],
                en_items[en_start:en_end],
                zh_items[zh_start:zh_end],
                en_docx,
                zh_docx,
                source_chapter=source_chapter,
                next_source_page=next_source_page,
                next_target_page=next_page,
            )
        lessons.append(lesson)

    return {
        "manual": "CSSC Six Sigma Black Belt Training Manual",
        "version": "0.2.0",
        "pageCount": 449,
        "chapters": lessons,
        "dictionary": CURATED_TERMS,
    }


def write_outputs(repo_root: Path, manual: dict[str, Any], en_docx: Path, zh_docx: Path) -> None:
    processed_dir = repo_root / "content" / "processed"
    chapters_dir = processed_dir / "chapters"
    chapters_dir.mkdir(parents=True, exist_ok=True)

    for chapter in manual["chapters"]:
        write_json(chapters_dir / f"{chapter['id']}.json", chapter)

    manifest = {
        "manual": manual["manual"],
        "version": manual["version"],
        "pageCount": manual["pageCount"],
        "chapters": [
            {
                "id": chapter["id"],
                "chapter": chapter["chapter"],
                "title": chapter["title"],
                "pageStart": chapter["pageStart"],
                "pageEnd": chapter["pageEnd"],
                "path": f"chapters/{chapter['id']}.json",
            }
            for chapter in manual["chapters"]
        ],
        "dictionary": "dictionary/six-sigma-terms.json",
    }
    write_json(processed_dir / "dictionary" / "six-sigma-terms.json", manual["dictionary"])
    write_json(processed_dir / "manifest.json", manifest)
    write_json(processed_dir / "manual.json", manual)

    generated_dir = repo_root / "apps" / "reader" / "src" / "generated"
    generated_dir.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(processed_dir / "dictionary" / "six-sigma-terms.json", generated_dir / "six-sigma-terms.json")

    public_content_dir = repo_root / "apps" / "reader" / "public" / "content"
    public_content_dir.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(processed_dir / "manual.json", public_content_dir / "manual.json")
    write_asset_files(repo_root, manual, en_docx, zh_docx)


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract the full bilingual manual into app JSON.")
    parser.add_argument("--en-docx", type=Path, default=DEFAULT_EN_DOCX)
    parser.add_argument("--zh-docx", type=Path, default=DEFAULT_ZH_DOCX)
    parser.add_argument("--repo-root", type=Path, default=DEFAULT_REPO_ROOT)
    args = parser.parse_args()
    manual = build_manual(args.en_docx, args.zh_docx, args.repo_root)
    write_outputs(args.repo_root, manual, args.en_docx, args.zh_docx)
    chapters = manual["chapters"]
    en_blocks = sum(len(section["content"]["en"]) for lesson in chapters for section in lesson["sections"])
    zh_blocks = sum(len(section["content"]["zh"]) for lesson in chapters for section in lesson["sections"])
    print(f"ok: extracted {len(chapters)} chapters, {en_blocks} English blocks, {zh_blocks} Chinese blocks")


if __name__ == "__main__":
    main()
